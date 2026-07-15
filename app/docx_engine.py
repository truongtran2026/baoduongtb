"""Fills [[placeholder]] tags in the legacy .docx templates and returns the
result as bytes (the database is the source of truth for every deployment,
including a serverless one with no persistent local disk). On a deployment
that does have a real local filesystem (the LAN app), the same bytes are
also best-effort mirrored to disk so "Mở thư mục" / Explorer still works
there - and that write goes through a path that can't hit Windows' MAX_PATH
crash.

Word frequently splits a single visible token like "[[thietbi]]" across
several <w:r> runs (spell-check, autocorrect, mid-edit formatting changes),
so a plain per-run string replace silently misses tags. The replacement
below works against each paragraph's *joined* text, then writes the result
back into the first run of the matched span and blanks out the rest -
matching what Word's own Find/Replace does visually.
"""
import hashlib
import io
import os
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document

PLACEHOLDER_RE = re.compile(r"\[\[(\w+)\]\]")
_ILLEGAL_CHARS_RE = re.compile(r'[\\/:*?"<>|\x00-\x1f]')
MAX_COMPONENT_LEN = 80  # keeps filenames readable in Explorer even though long paths are supported


def _iter_run_spans(runs):
    pos = 0
    for r in runs:
        length = len(r.text)
        yield r, pos, pos + length
        pos += length


def replace_placeholders_in_paragraph(paragraph, mapping: dict) -> int:
    """Replace every [[key]] in a paragraph with mapping[key]. Returns count replaced."""
    replaced = 0
    guard = 0
    while guard < 200:
        guard += 1
        runs = paragraph.runs
        if not runs:
            return replaced
        full_text = "".join(r.text for r in runs)
        match = PLACEHOLDER_RE.search(full_text)
        if not match:
            return replaced
        key = match.group(1)
        start, end = match.span()
        if key not in mapping:
            return replaced  # unknown tag: leave as literal text, stop scanning this paragraph

        value = str(mapping[key])
        touched = [(r, rs, re_) for r, rs, re_ in _iter_run_spans(runs) if re_ > start and rs < end]
        first_run, first_start, _ = touched[0]
        last_run, last_start, _ = touched[-1]

        before = first_run.text[: start - first_start]
        after = last_run.text[end - last_start :]
        if first_run is last_run:
            first_run.text = before + value + after
        else:
            first_run.text = before + value
            last_run.text = after
            for r, _, _ in touched[1:-1]:
                r.text = ""
        replaced += 1
    return replaced


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_all_paragraphs(document: Document):
    yield from document.paragraphs
    for t in document.tables:
        yield from _iter_table_paragraphs(t)
    for section in document.sections:
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            try:
                yield from part.paragraphs
                for t in part.tables:
                    yield from _iter_table_paragraphs(t)
            except Exception:
                continue


def find_placeholders(document: Document) -> set[str]:
    found = set()
    for p in _iter_all_paragraphs(document):
        found.update(PLACEHOLDER_RE.findall(p.text))
    return found


def fill_template(template_bytes: bytes, mapping: dict) -> Document:
    document = Document(io.BytesIO(template_bytes))
    for paragraph in _iter_all_paragraphs(document):
        replace_placeholders_in_paragraph(paragraph, mapping)
    return document


def render_to_bytes(template_bytes: bytes, mapping: dict) -> bytes:
    document = fill_template(template_bytes, mapping)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def sanitize_filename_component(text: str) -> str:
    text = _ILLEGAL_CHARS_RE.sub("_", text or "")
    text = re.sub(r"\s+", " ", text).strip()
    text = text.strip(". ") or "khong_ten"
    return text


def shorten_component(text: str, max_len: int = MAX_COMPONENT_LEN) -> tuple[str, bool]:
    if len(text) <= max_len:
        return text, False
    digest = hashlib.sha1(text.encode("utf-8")).hexdigest()[:6]
    keep = max(1, max_len - len(digest) - 1)
    return f"{text[:keep]}~{digest}", True


def build_output_filename(prefix: str, maintenance_date: date, device_name: str, station_code: str) -> tuple[str, bool]:
    device_part = sanitize_filename_component(device_name)
    station_part = sanitize_filename_component(station_code)
    device_part, truncated = shorten_component(device_part)
    filename = (
        f"{prefix}{maintenance_date:%m}-{maintenance_date:%Y%m%d}-{device_part}_{station_part}.docx"
    )
    return filename, truncated


def long_path(path: Path) -> str:
    """Return a save-path string that bypasses Windows' 260-char MAX_PATH limit.

    This is the direct fix for the legacy tool's crash: instead of hoping every
    generated name stays under MAX_PATH, we always opt into Windows' extended-length
    path form (\\\\?\\...), which supports paths up to ~32,000 characters.
    """
    resolved = str(path.resolve())
    if os.name == "nt" and not resolved.startswith("\\\\?\\"):
        return "\\\\?\\" + resolved
    return resolved


def ensure_dir(path: Path) -> None:
    os.makedirs(long_path(path), exist_ok=True)


WARN_PATH_LENGTH = 240  # heads-up threshold well under the classic 260-char limit


@dataclass
class GenerateResult:
    filename: str
    full_path: str
    content: bytes | None
    status: str  # success | warning | error
    message: str = ""


def generate_one(template_bytes: bytes, output_dir: Path | None, prefix: str, maintenance_date: date,
                  device_name: str, station_code: str, mapping: dict) -> GenerateResult:
    filename, truncated = build_output_filename(prefix, maintenance_date, device_name, station_code)

    warnings = []
    if truncated:
        warnings.append("Tên thiết bị quá dài nên đã được rút gọn trong tên file")

    try:
        content = render_to_bytes(template_bytes, mapping)
    except Exception as exc:  # noqa: BLE001 - surfaced to the user, one device shouldn't abort the whole batch
        return GenerateResult(filename=filename, full_path="", content=None, status="error", message=str(exc))

    full_path = ""
    if output_dir is not None:
        # Best-effort mirror to local disk (LAN deployment only - there's no
        # writable filesystem on a serverless deployment). The database copy
        # above is what downloads/history actually read from, so a failure
        # here never turns a successful generation into an error.
        target = output_dir / filename
        full_path = str(target)
        if len(full_path) > WARN_PATH_LENGTH:
            warnings.append("Đường dẫn khá dài (đã dùng chế độ chống lỗi Windows MAX_PATH)")
        try:
            ensure_dir(output_dir)
            Path(long_path(target)).write_bytes(content)
        except Exception:
            pass

    return GenerateResult(
        filename=filename,
        full_path=full_path,
        content=content,
        status="warning" if warnings else "success",
        message="; ".join(warnings),
    )
